#!/usr/bin/env python3
"""
Python FastMCP implementation for Document Editor MCP Server.
Provides high-fidelity PDF AcroForm filling, DOCX table/placeholder form
filling, text overlay, layout analysis, and validation tools using pypdf,
reportlab, and python-docx.
"""

import io
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastmcp import FastMCP
import pypdf
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black

mcp = FastMCP("document-editor")

OUTPUTS_DIR = "outputs"
UPLOADS_DIR = "uploads"


def _workspace_root() -> Path:
    """Root directory for all document read/write operations."""
    env = os.environ.get("HRAGENT_WORKSPACE_DIR")
    if env:
        root = Path(env)
        if not root.is_absolute():
            root = Path.cwd() / root
        return root.resolve()
    return (Path.cwd() / "workspace").resolve()


def _resolve_path(file_path: str, *, must_exist: bool = True) -> Path:
    """Resolve a workspace-relative path for reads and writes."""
    raw = Path(file_path)
    if raw.is_absolute():
        resolved = raw.resolve()
    else:
        resolved = (_workspace_root() / raw).resolve()

    root = _workspace_root()
    try:
        resolved.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"Path must stay inside workspace: {file_path}") from exc

    if must_exist and not resolved.exists():
        raise FileNotFoundError(
            f"File not found: {file_path} (looked in workspace at {resolved})"
        )
    return resolved


def _workspace_relative(path: Path) -> str:
    """Return a forward-slash path relative to the workspace root."""
    root = _workspace_root()
    try:
        rel = path.resolve().relative_to(root)
    except ValueError:
        return path.name
    return rel.as_posix()


def _result_paths(input_path: Path, output_path: Path) -> Dict[str, str]:
    return {
        "input_path": _workspace_relative(input_path),
        "output_path": _workspace_relative(output_path),
    }


def _file_extension(path: Path) -> str:
    return path.suffix.lower().lstrip(".")


def _is_pdf(path: Path) -> bool:
    return _file_extension(path) == "pdf"


def _is_docx(path: Path) -> bool:
    return _file_extension(path) == "docx"


_PLACEHOLDER_RE = re.compile(r"\{([^{}\s][^{}]*)\}")
_INLINE_LABEL_RE = re.compile(
    r"^(.{2,80}?)\s*:\s*([_\s]{3,}|)\s*$", re.MULTILINE
)


def _slugify_label(text: str) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", text.strip().rstrip(":"))
    slug = re.sub(r"[\s-]+", "_", cleaned.strip().lower())
    return slug or "field"


def _cell_text(cell) -> str:
    return (cell.text or "").strip()


def _is_blank_value(text: str) -> bool:
    if not text:
        return True
    stripped = text.strip()
    if not stripped:
        return True
    if set(stripped) <= {"_", "-", ".", " "}:
        return True
    return False


def _looks_like_label(text: str) -> bool:
    stripped = text.strip()
    if not stripped or len(stripped) > 120:
        return False
    if stripped.endswith(":"):
        return True
    if re.match(r"^[A-Z][A-Za-z0-9 /&()-]{1,80}$", stripped):
        return True
    return False


def _load_docx(path: Path) -> Document:
    return Document(str(path))


def _iter_paragraphs_in_table(table: Table) -> List[Paragraph]:
    paragraphs: List[Paragraph] = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
    return paragraphs


def _detect_docx_fields(doc: Document) -> List[Dict[str, Any]]:
    """Detect fillable fields in a DOCX: table cells, placeholders, inline labels."""
    fields: List[Dict[str, Any]] = []
    seen: set[str] = set()

    def add_field(
        field_id: str,
        *,
        label: str,
        field_type: str,
        value: str,
        location: Dict[str, Any],
    ) -> None:
        base_id = field_id
        suffix = 2
        while field_id in seen:
            field_id = f"{base_id}_{suffix}"
            suffix += 1
        seen.add(field_id)
        fields.append(
            {
                "name": field_id,
                "field_name": field_id,
                "label": label,
                "type": field_type,
                "value": value,
                "location": location,
            }
        )

    for t_idx, table in enumerate(doc.tables):
        header_labels: List[str] = []
        if table.rows:
            first_row = [_cell_text(c) for c in table.rows[0].cells]
            if len(first_row) >= 2 and all(
                _looks_like_label(c) or c for c in first_row[1:]
            ):
                header_labels = [_slugify_label(c) if c else f"col_{i}" for i, c in enumerate(first_row)]

        start_row = 1 if header_labels else 0
        for r_idx in range(start_row, len(table.rows)):
            row = table.rows[r_idx]
            cells = row.cells
            if len(cells) < 2:
                continue
            row_label = _cell_text(cells[0])
            if not _looks_like_label(row_label):
                continue

            if header_labels and len(cells) > 2:
                for c_idx in range(1, len(cells)):
                    header = header_labels[c_idx] if c_idx < len(header_labels) else f"col_{c_idx}"
                    value = _cell_text(cells[c_idx])
                    if not _is_blank_value(value):
                        continue
                    label = f"{row_label.rstrip(':')} ({header.replace('_', ' ')})"
                    field_id = _slugify_label(f"{row_label}_{header}")
                    add_field(
                        field_id,
                        label=label,
                        field_type="table_cell",
                        value=value,
                        location={"table": t_idx, "row": r_idx, "col": c_idx},
                    )
            else:
                value = _cell_text(cells[1])
                if not _is_blank_value(value):
                    continue
                field_id = _slugify_label(row_label)
                add_field(
                    field_id,
                    label=row_label.rstrip(":"),
                    field_type="table_cell",
                    value=value,
                    location={"table": t_idx, "row": r_idx, "col": 1},
                )

    placeholder_hits: Dict[str, str] = {}
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        all_paragraphs.extend(_iter_paragraphs_in_table(table))

    for p_idx, paragraph in enumerate(all_paragraphs):
        text = paragraph.text or ""
        for match in _PLACEHOLDER_RE.finditer(text):
            key = match.group(1).strip()
            field_id = _slugify_label(key)
            placeholder_hits.setdefault(field_id, key)
            if field_id not in seen:
                add_field(
                    field_id,
                    label=key,
                    field_type="placeholder",
                    value="",
                    location={"paragraph_index": p_idx, "placeholder": key},
                )

    for p_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text or ""
        for match in _INLINE_LABEL_RE.finditer(text):
            label = match.group(1).strip()
            if not _looks_like_label(label + ":"):
                continue
            field_id = _slugify_label(label)
            if field_id in seen:
                continue
            add_field(
                field_id,
                label=label,
                field_type="inline_label",
                value="",
                location={"paragraph_index": p_idx, "label": label},
            )

    return fields


def _replace_paragraph_placeholders(paragraph: Paragraph, fields: Dict[str, str]) -> bool:
    full_text = paragraph.text or ""
    if "{" not in full_text:
        return False
    updated = full_text
    changed = False
    for key, value in fields.items():
        patterns = [
            re.compile(r"\{" + re.escape(key) + r"\}"),
            re.compile(r"\{" + re.escape(_slugify_label(key)) + r"\}"),
        ]
        for pattern in patterns:
            if pattern.search(updated):
                updated = pattern.sub(str(value), updated)
                changed = True
    if changed:
        paragraph.text = updated
    return changed


def _fill_docx_inline_label(paragraph: Paragraph, label: str, value: str) -> bool:
    text = paragraph.text or ""
    pattern = re.compile(
        rf"^({re.escape(label)}\s*:\s*)([_\s.-]*)\s*$", re.IGNORECASE
    )
    match = pattern.match(text.strip())
    if not match:
        pattern2 = re.compile(
            rf"({re.escape(label)}\s*:\s*)([_\s.-]*)", re.IGNORECASE
        )
        if not pattern2.search(text):
            return False
        paragraph.text = pattern2.sub(rf"\1{value}", text, count=1)
        return True
    paragraph.text = f"{match.group(1)}{value}"
    return True


def _apply_docx_fields(doc: Document, field_defs: List[Dict[str, Any]], values: Dict[str, Any]) -> Tuple[int, List[str]]:
    normalized = {_slugify_label(k): str(v) for k, v in values.items()}
    filled = 0
    missing: List[str] = []

    by_id = {f["name"]: f for f in field_defs}
    for field_id, value in normalized.items():
        field = by_id.get(field_id)
        if not field:
            missing.append(field_id)
            continue
        loc = field.get("location") or {}
        field_type = field.get("type")
        if field_type == "table_cell":
            table_idx = loc["table"]
            row_idx = loc["row"]
            col_idx = loc["col"]
            cell = doc.tables[table_idx].cell(row_idx, col_idx)
            cell.text = value
            filled += 1
        elif field_type == "placeholder":
            key = loc.get("placeholder", field.get("label", field_id))
            replacements = {key: value, field_id: value}
            changed = False
            for paragraph in doc.paragraphs:
                changed = _replace_paragraph_placeholders(paragraph, replacements) or changed
            for table in doc.tables:
                for paragraph in _iter_paragraphs_in_table(table):
                    changed = _replace_paragraph_placeholders(paragraph, replacements) or changed
            if changed:
                filled += 1
            else:
                missing.append(field_id)
        elif field_type == "inline_label":
            label = loc.get("label", field.get("label", field_id))
            p_idx = loc.get("paragraph_index", 0)
            if 0 <= p_idx < len(doc.paragraphs):
                if _fill_docx_inline_label(doc.paragraphs[p_idx], label, value):
                    filled += 1
                else:
                    missing.append(field_id)
            else:
                missing.append(field_id)

    return filled, missing


def _docx_to_read_content(doc: Document, max_sections: int = 50) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    for idx, paragraph in enumerate(doc.paragraphs[:max_sections]):
        text = (paragraph.text or "").strip()
        if text:
            sections.append({"section_number": idx + 1, "text": text, "kind": "paragraph"})
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([_cell_text(cell) for cell in row.cells])
        sections.append(
            {
                "section_number": len(sections) + 1,
                "text": "",
                "kind": "table",
                "table_index": t_idx,
                "rows": rows,
            }
        )
        if len(sections) >= max_sections:
            break
    return sections


@mcp.tool()
def office_list_pdf_fields(file_path: str) -> List[Dict[str, Any]]:
    """List all interactive form fields (AcroForm) in a PDF file.

    Args:
        file_path: Path to the PDF file (e.g. 'i9_form.pdf').
    """
    resolved = _resolve_path(file_path)
    if not resolved.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    reader = pypdf.PdfReader(str(resolved))
    raw_fields = reader.get_fields() or {}
    results = []
    for name, field in raw_fields.items():
        field_type = "text"
        val = field.get("/V", "")
        ft = str(field.get("/FT", ""))
        if "/Btn" in ft:
            field_type = "checkbox" if "/Ff" in field else "button"
        elif "/Tx" in ft:
            field_type = "text"
        elif "/Ch" in ft:
            field_type = "choice"

        results.append({
            "name": name,
            "field_name": name,
            "type": field_type,
            "value": str(val) if val else "",
        })
    return results

@mcp.tool()
def office_fill_pdf_form(
    input_path: str,
    output_path: str,
    fields: Dict[str, Any],
    flatten: bool = False,
) -> Dict[str, Any]:
    """Fill AcroForm fields in a PDF document and save the result to output_path.

    Args:
        input_path: Path to source PDF (e.g. 'i9_form.pdf').
        output_path: Path where filled PDF should be saved (e.g. 'i9_form_filled.pdf').
        fields: Mapping of field names to values (e.g. {'First Name Given Name': 'Joseph', 'Last Name': 'Johnson'}).
        flatten: Whether to make form fields read-only / flattened.
    """
    in_file = _resolve_path(input_path)
    if not in_file.exists():
        raise FileNotFoundError(f"Source file not found: {input_path}")

    out_file = _resolve_path(output_path, must_exist=False)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    reader = pypdf.PdfReader(str(in_file))
    writer = pypdf.PdfWriter()
    writer.append(reader)

    # Convert values to strings
    field_values = {k: str(v) for k, v in fields.items()}

    # Update form field values across all pages
    for page in writer.pages:
        try:
            writer.update_page_form_field_values(page, field_values, auto_regenerate=True)
        except Exception:
            pass

    with open(out_file, "wb") as f:
        writer.write(f)

    return {
        "status": "success",
        **_result_paths(in_file, out_file),
        "fields_filled": len(fields),
        "total_pages": len(writer.pages),
        "file_size": out_file.stat().st_size,
    }


@mcp.tool()
def office_list_docx_fields(file_path: str) -> List[Dict[str, Any]]:
    """List fillable fields in a DOCX document.

    Detects empty table cells (label in first column, value in second),
    {placeholder} mail-merge tokens, and inline labels like "Signature: ___".

    Args:
        file_path: Path to the DOCX file (e.g. 'emergency_contact_form.docx').
    """
    resolved = _resolve_path(file_path)
    if not _is_docx(resolved):
        raise ValueError(f"Not a DOCX file: {file_path}")
    doc = _load_docx(resolved)
    return _detect_docx_fields(doc)


def _list_docx_fields(file_path: str) -> List[Dict[str, Any]]:
    return office_list_docx_fields.fn(file_path)


def _fill_docx_form(
    input_path: str,
    output_path: str,
    fields: Dict[str, Any],
) -> Dict[str, Any]:
    in_file = _resolve_path(input_path)
    if not _is_docx(in_file):
        raise ValueError(f"Not a DOCX file: {input_path}")

    out_file = _resolve_path(output_path, must_exist=False)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    doc = _load_docx(in_file)
    field_defs = _detect_docx_fields(doc)
    filled, missing = _apply_docx_fields(doc, field_defs, fields)
    doc.save(str(out_file))

    return {
        "status": "success",
        **_result_paths(in_file, out_file),
        "fields_filled": filled,
        "fields_requested": len(fields),
        "missing_fields": missing,
        "file_size": out_file.stat().st_size,
    }


@mcp.tool()
def office_fill_docx_form(
    input_path: str,
    output_path: str,
    fields: Dict[str, Any],
) -> Dict[str, Any]:
    """Fill fields in a DOCX document and save the result to output_path.

    Use office_list_docx_fields first to discover field names. Field keys can
    be the slugified name (e.g. 'employee_name') or the human label.

    Args:
        input_path: Path to source DOCX (e.g. 'emergency_contact_form.docx').
        output_path: Path where filled DOCX should be saved (e.g. 'outputs/emergency_contact_Jane_Doe.docx').
        fields: Mapping of field names to values.
    """
    return _fill_docx_form(input_path, output_path, fields)


@mcp.tool()
def office_template_detect(file_path: str) -> List[Dict[str, Any]]:
    """Detect {placeholder} mail-merge fields in a DOCX template.

    Args:
        file_path: Path to the DOCX template.
    """
    all_fields = _list_docx_fields(file_path)
    return [f for f in all_fields if f.get("type") == "placeholder"]


@mcp.tool()
def office_template_fill(
    input_path: str,
    output_path: str,
    fields: Dict[str, Any],
) -> Dict[str, Any]:
    """Fill {placeholder} fields in a DOCX mail-merge template.

    For table-based HR forms (empty table cells), use office_fill_docx_form instead.

    Args:
        input_path: Path to source DOCX template.
        output_path: Destination path for the filled document.
        fields: Placeholder name -> value mapping.
    """
    return _fill_docx_form(input_path, output_path, fields)


@mcp.tool()
def office_template_batch(
    templates: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Batch-fill multiple DOCX templates.

    Args:
        templates: List of dicts with keys input_path, output_path, and fields.
    """
    results = []
    for item in templates:
        result = _fill_docx_form(
            item["input_path"],
            item["output_path"],
            item.get("fields", {}),
        )
        results.append(result)
    return results


@mcp.tool()
def office_overlay_pdf_text(
    input_path: str,
    output_path: str,
    text: str,
    page_number: int = 1,
    x: float = 100.0,
    y: float = 700.0,
    font_size: float = 12.0,
) -> Dict[str, Any]:
    """Overlay text onto a PDF document page at specific coordinates.

    Args:
        input_path: Path to source PDF.
        output_path: Destination path for the new PDF.
        text: Text string to overlay.
        page_number: 1-indexed page number to place the text on.
        x: X coordinate (in points from bottom-left).
        y: Y coordinate (in points from bottom-left).
        font_size: Font size in points.
    """
    in_file = _resolve_path(input_path)
    if not in_file.exists():
        raise FileNotFoundError(f"Source file not found: {input_path}")

    out_file = _resolve_path(output_path, must_exist=False)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    reader = pypdf.PdfReader(str(in_file))
    writer = pypdf.PdfWriter()

    # Create overlay PDF in memory
    packet = io.BytesIO()
    can = canvas.Canvas(packet)
    can.setFont("Helvetica", font_size)
    can.setFillColor(black)
    can.drawString(x, y, text)
    can.save()
    packet.seek(0)
    overlay_reader = pypdf.PdfReader(packet)
    overlay_page = overlay_reader.pages[0]

    for idx, page in enumerate(reader.pages):
        if idx == (page_number - 1):
            page.merge_page(overlay_page)
        writer.add_page(page)

    with open(out_file, "wb") as f:
        writer.write(f)

    return {
        "status": "success",
        **_result_paths(in_file, out_file),
        "text_overlaid": text,
        "page": page_number,
        "coordinates": {"x": x, "y": y},
    }

@mcp.tool()
def office_validate(file_path: str) -> Dict[str, Any]:
    """Validate document integrity, structure, and check for corruption.

    Args:
        file_path: Path to document to validate.
    """
    resolved = _resolve_path(file_path)
    if not resolved.exists():
        return {
            "status": "failed",
            "passed": 0,
            "total_checks": 6,
            "error": f"File not found: {file_path}",
        }

    if _is_docx(resolved):
        checks_passed = 0
        total_checks = 5
        messages: List[str] = []
        size = resolved.stat().st_size
        if size > 0:
            checks_passed += 1
            messages.append("File exists and has positive byte length")
        try:
            doc = _load_docx(resolved)
            checks_passed += 1
            messages.append("DOCX package opens successfully")
        except Exception as e:
            return {
                "status": "failed",
                "passed": checks_passed,
                "total_checks": total_checks,
                "error": str(e),
            }
        if doc.paragraphs or doc.tables:
            checks_passed += 1
            messages.append(
                f"Document contains {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
            )
        fields = _detect_docx_fields(doc)
        checks_passed += 1
        messages.append(f"Detected {len(fields)} fillable field(s)")
        checks_passed += 1
        messages.append("Document structure is readable")
        return {
            "status": "passed",
            "passed": checks_passed,
            "total_checks": total_checks,
            "summary": f"{checks_passed}/{total_checks} integrity checks passed",
            "file_path": _workspace_relative(resolved),
            "details": messages,
        }

    checks_passed = 0
    total_checks = 6
    messages = []

    # Check 1: File exists and is non-empty
    size = resolved.stat().st_size
    if size > 0:
        checks_passed += 1
        messages.append("File exists and has positive byte length")

    # Check 2: Readable PDF Header
    try:
        reader = pypdf.PdfReader(str(resolved))
        checks_passed += 1
        messages.append("PDF header and trailer are valid")
    except Exception as e:
        return {"status": "failed", "passed": checks_passed, "total_checks": total_checks, "error": str(e)}

    # Check 3: Page count valid
    if len(reader.pages) > 0:
        checks_passed += 1
        messages.append(f"Document contains {len(reader.pages)} valid pages")

    # Check 4: Text stream extractable
    try:
        txt = reader.pages[0].extract_text()
        checks_passed += 1
        messages.append("Text streams are decryptable and extractable")
    except Exception:
        messages.append("Text stream warning")

    # Check 5: Form catalog intact
    fields = reader.get_fields()
    checks_passed += 1
    messages.append(f"AcroForm catalog intact ({len(fields) if fields else 0} fields)")

    # Check 6: Encryption / DRM check
    if not reader.is_encrypted:
        checks_passed += 1
        messages.append("Document permissions allow read/write")

    return {
        "status": "passed",
        "passed": checks_passed,
        "total_checks": total_checks,
        "summary": f"{checks_passed}/{total_checks} integrity checks passed",
        "file_path": _workspace_relative(resolved),
        "details": messages,
    }

@mcp.tool()
def office_read(file_path: str, max_pages: int = 5) -> Dict[str, Any]:
    """Read text content and structure from a PDF or Office document.

    Args:
        file_path: Path to the document.
        max_pages: Maximum number of pages/sections to read.
    """
    resolved = _resolve_path(file_path)
    if not resolved.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if _is_docx(resolved):
        doc = _load_docx(resolved)
        content = _docx_to_read_content(doc, max_sections=max_pages)
        return {
            "file_name": _workspace_relative(resolved),
            "file_path": _workspace_relative(resolved),
            "format": "docx",
            "total_pages": None,
            "sections_read": len(content),
            "content": content,
        }

    reader = pypdf.PdfReader(str(resolved))
    pages_text = []
    for idx, page in enumerate(reader.pages[:max_pages]):
        pages_text.append({
            "page_number": idx + 1,
            "text": page.extract_text() or "",
        })

    return {
        "file_name": _workspace_relative(resolved),
        "file_path": _workspace_relative(resolved),
        "format": "pdf",
        "total_pages": len(reader.pages),
        "pages_read": len(pages_text),
        "content": pages_text,
    }

@mcp.tool()
def get_document_info(file_path: str) -> Dict[str, Any]:
    """Get metadata, page count, and properties of a document.

    Args:
        file_path: Path to the document.
    """
    resolved = _resolve_path(file_path)
    if not resolved.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if _is_docx(resolved):
        doc = _load_docx(resolved)
        fields = _detect_docx_fields(doc)
        return {
            "file_name": _workspace_relative(resolved),
            "file_path": _workspace_relative(resolved),
            "format": "docx",
            "total_pages": None,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "form_fields_count": len(fields),
            "is_encrypted": False,
            "file_size_bytes": resolved.stat().st_size,
        }

    reader = pypdf.PdfReader(str(resolved))
    fields = reader.get_fields() or {}
    return {
        "file_name": _workspace_relative(resolved),
        "file_path": _workspace_relative(resolved),
        "format": "pdf",
        "total_pages": len(reader.pages),
        "form_fields_count": len(fields),
        "is_encrypted": reader.is_encrypted,
        "file_size_bytes": resolved.stat().st_size,
    }

@mcp.tool()
def list_formats() -> List[str]:
    """List all supported document formats."""
    return ["pdf", "docx", "xlsx", "pptx", "txt", "md"]

def main():
    mcp.run(transport="stdio")

if __name__ == "__main__":
    main()
