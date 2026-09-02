"""Tests for DOCX form listing and filling in document_server."""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document


_SERVER_PATH = (
    Path(__file__).resolve().parents[1]
    / "marketplaces"
    / "integrations"
    / "document-editor"
    / "server"
    / "document_server.py"
)


def _load_document_server():
    spec = importlib.util.spec_from_file_location("document_server", _SERVER_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _make_table_form(path: Path) -> None:
    doc = Document()
    doc.add_heading("Emergency Contact Form", level=1)
    table = doc.add_table(rows=3, cols=2)
    rows = [
        ("Employee Name", ""),
        ("Employee ID", ""),
        ("Department", ""),
    ]
    for idx, (label, value) in enumerate(rows):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value
    doc.save(str(path))


def _make_placeholder_form(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Offer letter for {employee_name}.")
    doc.add_paragraph("Start date: {start_date}.")
    doc.save(str(path))


def _make_inline_form(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Code of Conduct Acknowledgment")
    doc.add_paragraph("Employee Name: _______________")
    doc.add_paragraph("Signature: _______________")
    doc.add_paragraph("Date: _______________")
    doc.save(str(path))


@pytest.fixture
def doc_server(tmp_path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    _make_table_form(workspace / "emergency_contact_form.docx")
    _make_placeholder_form(workspace / "offer_letter.docx")
    _make_inline_form(workspace / "code_of_conduct_acknowledgment.docx")
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("HRAGENT_WORKSPACE_DIR", str(workspace))
    return _load_document_server()


def _tool(module, name: str):
    """Call an MCP-wrapped tool's underlying function."""
    fn = getattr(module, name)
    return fn.fn if hasattr(fn, "fn") else fn

def test_list_docx_table_fields(doc_server):
    fields = _tool(doc_server, "office_list_docx_fields")("emergency_contact_form.docx")
    names = {f["name"] for f in fields}
    assert "employee_name" in names
    assert "employee_id" in names
    assert "department" in names
    assert all(f["type"] == "table_cell" for f in fields)


def test_fill_docx_table_form(doc_server):
    result = _tool(doc_server, "office_fill_docx_form")(
        "emergency_contact_form.docx",
        "outputs/emergency_contact_filled.docx",
        {
            "employee_name": "Jane Doe",
            "employee_id": "E-1001",
            "department": "Engineering",
        },
    )
    assert result["status"] == "success"
    assert result["fields_filled"] == 3

    out = doc_server._resolve_path("outputs/emergency_contact_filled.docx")
    doc = Document(str(out))
    table = doc.tables[0]
    assert table.rows[0].cells[1].text == "Jane Doe"
    assert table.rows[1].cells[1].text == "E-1001"
    assert table.rows[2].cells[1].text == "Engineering"


def test_list_and_fill_placeholder_template(doc_server):
    placeholders = _tool(doc_server, "office_template_detect")("offer_letter.docx")
    names = {p["name"] for p in placeholders}
    assert "employee_name" in names
    assert "start_date" in names

    result = _tool(doc_server, "office_template_fill")(
        "offer_letter.docx",
        "outputs/offer_letter_filled.docx",
        {"employee_name": "Alex Kim", "start_date": "2026-09-01"},
    )
    assert result["fields_filled"] == 2
    out = doc_server._resolve_path("outputs/offer_letter_filled.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Alex Kim" in text
    assert "2026-09-01" in text
    assert "{employee_name}" not in text


def test_fill_inline_signature_fields(doc_server):
    fields = _tool(doc_server, "office_list_docx_fields")("code_of_conduct_acknowledgment.docx")
    names = {f["name"] for f in fields}
    assert "employee_name" in names
    assert "signature" in names
    assert "date" in names

    result = _tool(doc_server, "office_fill_docx_form")(
        "code_of_conduct_acknowledgment.docx",
        "outputs/coc_filled.docx",
        {
            "employee_name": "Jane Doe",
            "signature": "Jane Doe",
            "date": "2026-08-19",
        },
    )
    assert result["fields_filled"] == 3
    out = doc_server._resolve_path("outputs/coc_filled.docx")
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Employee Name: Jane Doe" in text
    assert "Signature: Jane Doe" in text
    assert "Date: 2026-08-19" in text


def test_office_read_docx(doc_server):
    result = _tool(doc_server, "office_read")("emergency_contact_form.docx")
    assert result["format"] == "docx"
    assert any(item.get("kind") == "table" for item in result["content"])


def test_office_validate_docx(doc_server):
    result = _tool(doc_server, "office_validate")("emergency_contact_form.docx")
    assert result["status"] == "passed"
    assert result["passed"] == result["total_checks"]
