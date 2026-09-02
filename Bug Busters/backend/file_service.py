from io import BytesIO
from html import escape

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.colors import HexColor, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable


# Bug Busters brand palette (matches the frontend's design tokens).
INK = HexColor("#241a45")
LILAC = HexColor("#b9a6ff")
ACCENT = HexColor("#6b4ce6")
ACCENT_700 = HexColor("#4a2bb5")
TEXT = HexColor("#1d1633")
TEXT_MUTED = HexColor("#6b6280")

HEADER_HEIGHT = 0.62 * 72  # points


def clean_pdf_text(text: str) -> str:
    """
    Replace Unicode characters that may not render correctly
    with ReportLab's default fonts.
    """
    replacements = {
        "‐": "-",   # hyphen
        "‑": "-",   # non-breaking hyphen
        "‒": "-",   # figure dash
        "–": "-",   # en dash
        "—": "-",   # em dash
        "−": "-",   # minus sign
        " ": " ",   # non-breaking space
        "‘": "'",   # left single quote
        "’": "'",   # right single quote
        "“": '"',   # left double quote
        "”": '"',   # right double quote
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text


def _split_heading_and_body(block: str) -> tuple[str, str]:
    """
    Generated policy text is consistently "Heading\nBody..." per
    blank-line-separated block (both the local templates and the AI
    generator follow this shape). Split on the first newline so the
    heading can be styled distinctly from its body.
    """
    if "\n" in block:
        heading, _, body = block.partition("\n")
        return heading.strip(), body.strip()

    return block.strip(), ""


def policy_to_docx_bytes(policy_content: str, title: str) -> bytes:
    doc = Document()

    title_paragraph = doc.add_heading(level=0)
    title_run = title_paragraph.add_run(title)
    title_run.font.color.rgb = RGBColor(0x24, 0x1A, 0x45)
    title_run.font.size = Pt(26)

    kicker = doc.add_paragraph()
    kicker_run = kicker.add_run("AI POLICY GENERATOR · BUG BUSTERS")
    kicker_run.font.size = Pt(9)
    kicker_run.font.bold = True
    kicker_run.font.color.rgb = RGBColor(0x6B, 0x4C, 0xE6)
    kicker.paragraph_format.space_after = Pt(18)

    for block in policy_content.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        heading, body = _split_heading_and_body(block)

        heading_paragraph = doc.add_paragraph()
        heading_run = heading_paragraph.add_run(heading)
        heading_run.font.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(0x4A, 0x2B, 0xB5)
        heading_paragraph.paragraph_format.space_before = Pt(14)
        heading_paragraph.paragraph_format.space_after = Pt(4)

        if body:
            body_paragraph = doc.add_paragraph(body)
            body_paragraph.paragraph_format.space_after = Pt(6)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()


def _draw_header_band(canvas, doc):
    canvas.saveState()

    page_width, page_height = doc.pagesize

    canvas.setFillColor(INK)
    canvas.rect(0, page_height - HEADER_HEIGHT, page_width, HEADER_HEIGHT, fill=1, stroke=0)

    canvas.setFillColor(white)
    canvas.setFont("Helvetica-Bold", 14)
    canvas.drawString(0.75 * 72, page_height - HEADER_HEIGHT / 2 - 5, "Bug Busters")

    canvas.setFillColor(LILAC)
    canvas.setFont("Helvetica", 8)
    canvas.drawString(2.05 * 72, page_height - HEADER_HEIGHT / 2 - 4, "AI POLICY GENERATOR")

    canvas.setFillColor(TEXT_MUTED)
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(page_width - 0.75 * 72, 0.4 * 72, f"Page {doc.page}")

    canvas.setStrokeColor(ACCENT)
    canvas.setLineWidth(2)
    canvas.line(0.75 * 72, 0.55 * 72, page_width - 0.75 * 72, 0.55 * 72)

    canvas.restoreState()


def policy_to_pdf_bytes(policy_content: str, title: str) -> bytes:
    buffer = BytesIO()

    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        title=title,
        topMargin=HEADER_HEIGHT + 28,
        bottomMargin=0.85 * 72,
        leftMargin=0.75 * 72,
        rightMargin=0.75 * 72,
    )

    title_style = ParagraphStyle(
        "PolicyTitle",
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=28,
        textColor=INK,
        spaceAfter=4,
    )

    heading_style = ParagraphStyle(
        "SectionHeading",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=ACCENT_700,
        spaceBefore=16,
        spaceAfter=4,
    )

    body_style = ParagraphStyle(
        "PolicyBody",
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=TEXT,
        alignment=TA_LEFT,
    )

    story = []

    story.append(Paragraph(escape(clean_pdf_text(title)), title_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=14))

    for block in policy_content.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        heading, body = _split_heading_and_body(block)

        story.append(Paragraph(escape(clean_pdf_text(heading)), heading_style))

        if body:
            safe_body = escape(clean_pdf_text(body)).replace("\n", "<br/>")
            story.append(Paragraph(safe_body, body_style))
            story.append(Spacer(1, 4))

    pdf.build(story, onFirstPage=_draw_header_band, onLaterPages=_draw_header_band)

    buffer.seek(0)

    return buffer.read()
